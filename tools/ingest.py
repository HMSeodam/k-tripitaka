#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
藏經 대조본 인게스트 파이프라인
------------------------------------------------------------------
입력:  sources/<work_id>/원문.txt        (CBETA/SAT/KABC 복사본, 선택)
       sources/<work_id>/번역.docx       (한국어 대조 번역, 선택)
       sources/works.yml 대신 tools/registry.json 에 서지 메타 기술
출력:  data/works/<work_id>.json         (열람·검색용 정규화 단위)
       data/manifest.json                (문헌 목록 + 통계)

핵심 설계
  1) 정렬 키는 '위치표지'(CBETA 行標, 예: [0287b17]) 이다.
     CBETA/SAT/KABC 가 공유하는 유일한 안정 좌표이므로 이것을 단위 ID로 삼는다.
  2) docx 는 문서마다 스타일 이름이 제각각이므로(Source Text / Chinese Source /
     대조 원문 / Normal ...) 스타일에 의존하지 않는다. 대신 '문자 스크립트 비율'로
     한문 단락과 한글 단락을 판정한다. 어떤 형식의 번역 docx 가 들어와도 동작한다.
  3) 위치표지가 없는 docx 는 한문 단락의 앞머리를 원문 TXT 와 대조하여 정렬한다.
"""

import json, os, re, sys, unicodedata, hashlib
from pathlib import Path

try:
    import docx  # python-docx
except ImportError:
    docx = None

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "sources"
OUT = ROOT / "data"

# ── 위치표지: [0287b17] / [1089a06] 형태 ──────────────────────────────
RE_MARKER = re.compile(r"\[(\d{3,4}[abc]\d{2})\]")
# 교감 표지: [1] [＊] [12] 등 (본문 아님)
RE_APPARATUS = re.compile(r"\[(?:\d{1,3}|＊|\*)\]")
RE_CJK = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
RE_HANGUL = re.compile(r"[\uac00-\ud7a3]")

# 번역 docx 에서 반복되는 '라벨' 단락 (내용 아님)
LABELS = {
    "원문", "한문 원문", "한국어 번역", "한국어 직역", "번역", "직역",
    "【註】", "【本文】", "【科文】", "【科層】",
    "【논서 번역】", "【주석서 번역】", "【번역 각주】", "번역 각주",
}
# ── 표제(heading) 분류 ────────────────────────────────────────────────
# 편집 안내·용어표·검증 보고 등 '해제' 성격의 표제. 이 아래의 한국어 문단은
# 번역이 아니므로 본문 단위에 붙이지 않는다.
RE_APPARATUS_HEAD = re.compile(
    r"^[A-Z]\.\s|"
    r"문헌\s*정보|편집|안내|범례|용어\s*(?:표|대응)|검증|검수|선독|판본\s*정보|"
    r"교정\s*(?:기록|보고|사항)|정리\s*(?:기록|보고|사항)|작업\s*(?:기록|보고)|"
    r"(?:정리|편집|번역|작업|교열|대조|처리|표기|배열)\s*(?:기준|방침|원칙|규칙|방법)|"
    r"(?:최종|전체|일괄)\s*(?:정리|교열|점검|확인)|"
    r"(?:편집|번역|대조|원문)\s*(?:검증|검수|확인)\s*(?:요약|결과|보고)?|"
    r"검증\s*요약|검수\s*요약|"
    r"번역상|쟁점|한계|읽는\s*법|과단\s*계층|부록|참고\s*문헌|대역\s*목차"
)
# 「1. …」 「2.3 …」 처럼 번호를 매긴 표제는 번역 docx 가 스스로 붙인
# 문서 차례이지 저본의 권·품이 아니다. 저본의 표제는 「제1 발심」·「권제1」·
# 「初篇」 꼴이라 이 규칙에 걸리지 않는다.
RE_DOC_OUTLINE = re.compile(r"^\d{1,2}(?:\.\d{1,2})*[.、]\s|^\d{1,2}(?:\.\d{1,2})+\s")
# 단위 라벨: 위치표지·문단 번호·「원문/번역」 같은 꼬리표만으로 이루어진 줄.
# 한글이 섞여 있어 번역문으로 오인되기 쉬우므로 본문에 들어가기 전에 걸러 낸다.
#   예) "[1206c22]  문단 001 · 원문" / "문단 003 · 한국어 직역" / "005. 번역"
RE_UNIT_LABEL = re.compile(
    r"^\s*(?:\[\d{3,4}[abc]\d{2}\])?\s*"
    r"(?:(?:문단|단락|단위|제)\s*\d+\s*(?:항|번|단)?)?\s*"
    r"(?:\d{1,4}\s*[.)]?)?\s*"
    r"(?:[·・|/:：\-—~]\s*)?"
    r"(?:한문\s*)?(?:한국어\s*)?"
    r"(?:원문|본문|번역|직역|대조|원문·번역)?\s*$"
)
LABEL_WORDS = ("원문", "번역", "직역", "본문", "대조", "문단", "단락", "단위")

# 단락 앞머리에 붙는 꼬리표: "원문  大乘起信論…" / "번역  『대승기신론소』…"
# 이 두 글자 때문에 짧은 한문 줄이 번역으로 오인되므로 먼저 떼어 낸다.
RE_SIDE_PREFIX = re.compile(
    r"^(원문|한문\s*원문|본문|번역|직역|한국어\s*직역|한국어\s*번역)\s*[:：·|]?\s+"
)


def strip_side(t: str):
    """앞머리 꼬리표를 떼고, 그것이 가리키던 쪽('cn'/'ko'/None)을 함께 돌려준다."""
    m = RE_SIDE_PREFIX.match(t)
    if not m:
        return t, None
    head = re.sub(r"\s", "", m.group(1))
    side = "cn" if head in ("원문", "한문원문", "본문") else "ko"
    return t[m.end():].strip(), side


def is_unit_label(t: str) -> bool:
    """내용 없는 단위 라벨인가. 실제 번역문을 잘못 지우지 않도록 조건을 좁게 둔다."""
    t = t.strip()
    if not t or len(t) > 40:
        return False
    if not any(w in t for w in LABEL_WORDS):
        return False
    return bool(RE_UNIT_LABEL.match(t))


# 개별 단위마다 붙는 표제(구조가 아님)
RE_UNIT_HEAD = re.compile(
    r"^(?:문단\s*\d+|\d{2,4}\.\s|권두\s*표제|\[\d{3,4}[abc]\d{2}\])"
)
# 라벨 성격의 표제 — 무시하되 단위를 끊지 않는다
RE_NOTE_HEAD = re.compile(r"(?:^|\s)(?:주|주석|각주|교감)$")
# 각주 묶음의 칸 이름. '교감·번역 메모'처럼 스타일 구분 없이 라벨 문단만으로
# 각주 시작을 알리는 docx 가 있어, 이 줄 아래는 각주로 모은다.
# 각주 묶음의 칸 이름. 라벨의 '끝 낱말'이 각주 계열이어야 한다.
#   주석·교감 / 교감 / 교감·번역 메모  → 각주 칸
#   【주석서 번역】 / 교감 표시 규칙     → 각주 칸이 아님
RE_NOTE_LABEL_HEAD = re.compile(
    r"^(?=.{1,16}$)[^\[\]\n。.]*"
    r"(?:메모|주기|비고|각주|주석|교감|교정)"
    r"\s*[】\]）)]?\s*[:：]?$")
# 다만 '교감 표시 규칙'처럼 안내 표제인 것은 각주 칸 이름이 아니다.
RE_NOTE_LABEL_NOT = re.compile(r"규칙|원칙|방식|기준|방법|안내|범례|일러두기")
# 칸 이름에 '· 계속' 같은 꼬리가 붙는 docx 가 있다.
# ('원문 · 계속' / '한국어 직역 · 계속') 꼬리를 떼고 칸 이름을 견준다.
RE_LABEL_TAIL = re.compile(
    r"\s*[·・|/／\-–—]\s*(?:계속|이어짐|이어서|continued|cont\.?)\s*$", re.I)


# 도판 자리를 말로 대신한 번역 문단. 도판을 띄우면 필요 없어진다.
RE_FIG_CAPTION = re.compile(
    r"(?:문자\s*본문|본문\s*문자)[^。.\n]{0,24}?(?:없|비어)"
    r"|이\s*(?:위치|자리)에[^。.\n]{0,24}?(?:도판|그림|圖|계도|界圖)"
    r"[^。.\n]{0,12}?(?:배치|삽입|들어)")


def label_key(t: str) -> str:
    return RE_LABEL_TAIL.sub("", t).rstrip("：: ").strip()


LABEL_HEADS = {"원문", "원문 목차", "한국어 목차", "한국어 직역", "한국어 번역",
               "번역", "직역"}
# 원문 안의 권 표제 (원문 TXT 만 있는 문헌의 분권 검출용)
RE_JUAN_LINE = re.compile(
    r"^.{0,24}?卷(?:第[一二三四五六七八九十百]+|[上中下])"
    r"(?:之[上中下]|[上中下])?(?:[(（].{0,12}?[)）])?\s*$"
)

NOTE_PREFIX = re.compile(
    r"^\s*(?:주석\s*\d*|각주\s*\d*|교감[·\s]|번역 각주|【번역 각주】|\[구두 불확실\]|\[판독 불확실\]|\[구문 불확실\])"
)


# ── 문자 판정 ────────────────────────────────────────────────────────
def hangul_ratio(s: str) -> float:
    letters = [c for c in s if unicodedata.category(c).startswith("L")]
    if not letters:
        return 0.0
    return sum(1 for c in letters if RE_HANGUL.match(c)) / len(letters)


def is_source_line(s: str) -> bool:
    """한문 원문 단락인가.

    번역문에는 반드시 한글이 섞이므로 한글 비율로 가른다.
    卷首·卷尾·諦 처럼 두세 글자짜리 원문 조각도 놓치지 않도록
    길이 조건은 최소한으로만 둔다."""
    body = RE_MARKER.sub("", s).strip()
    body = re.sub(r"^[【〔\[(（]+|[】〕\])）]+$", "", body).strip()
    if len(body) < 2:
        return False
    if not RE_CJK.search(body):
        return False
    return hangul_ratio(body) < 0.12


# ── 검색용 정규화(이체자 폴딩) ────────────────────────────────────────
# CBETA·SAT·KABC 판본 간 흔들리는 글자만 최소로 접는다. 필요할 때 추가하면 된다.
VARIANTS = {
    # ① 한국에서 쓰는 자형 → CBETA·SAT 판본의 자형
    #    (한글 글꼴로 입력한 검색어가 원문과 어긋나는 것을 막는다)
    "卽": "即", "敎": "教", "眞": "真", "爲": "為", "衆": "眾", "敍": "敘",
    "兪": "俞", "靑": "青", "絶": "絕", "戸": "戶", "户": "戶", "幷": "并",
    "擧": "舉", "恆": "恒", "冊": "册", "郞": "郎", "祕": "秘", "麤": "麁",
    "况": "況", "决": "決", "畧": "略", "凈": "淨", "裏": "裡",
    "鑒": "鑑", "慙": "慚", "飜": "翻", "遲": "遅", "旣": "既", "飮": "飲",
    "衞": "衛", "羣": "群", "菴": "庵", "翫": "玩", "皁": "皂", "牀": "床",
    "氷": "冰", "朶": "朵", "擡": "抬", "掛": "挂", "覩": "睹", "踰": "逾",
    "顚": "顛", "鬪": "鬥", "姊": "姉", "妬": "妒", "嶽": "岳", "疊": "叠",
    "甎": "磚", "筯": "箸", "禱": "祷", "稟": "禀", "剏": "創", "勑": "敕",
    "勅": "敕", "軆": "體", "釼": "劍", "劒": "劍", "閒": "間", "餧": "餵",
    "髣": "仿", "賔": "賓", "栰": "筏", "灋": "法", "噐": "器", "瑯": "琅",
    "逈": "迥", "盋": "缽", "鉢": "缽", "蘂": "蕊",

    # ② 일본 신자체 → 판본의 자형 (SAT 계열 자료를 함께 찾기 위함)
    "歴": "歷", "虚": "虛", "温": "溫", "説": "說", "増": "增", "徳": "德",
    "対": "對", "実": "實", "経": "經", "処": "處", "帰": "歸", "独": "獨",
    "変": "變", "数": "數", "関": "關", "両": "兩", "称": "稱", "証": "證",
    "釈": "釋", "寿": "壽", "薬": "藥", "礼": "禮", "発": "發", "国": "國",
    "気": "氣", "楽": "樂", "辞": "辭", "悪": "惡", "盗": "盜", "頼": "賴",
    "黄": "黃", "斉": "齊", "竜": "龍", "亀": "龜", "尽": "盡", "旧": "舊",
    "声": "聲", "触": "觸", "読": "讀", "鉄": "鐵", "雑": "雜", "霊": "靈",
    "黙": "默", "斎": "齋", "歯": "齒", "禅": "禪", "静": "靜", "黒": "黑",
    "摂": "攝", "継": "繼", "繊": "纖", "聡": "聰", "粛": "肅", "臈": "臘",
    "軽": "輕", "駆": "驅", "篭": "籠", "栄": "榮", "荘": "莊", "剣": "劍",
    "円": "圓", "戦": "戰", "帯": "帶", "広": "廣", "厳": "嚴", "覚": "覺",
    "観": "觀", "壊": "壞", "頬": "頰", "巻": "卷", "昼": "晝",
    "写": "寫", "壌": "壤", "収": "收", "拝": "拜", 
    # ③ 중국 간체 → 번체 (간체로 입력한 검색어도 함께 찾는다)
    "众": "眾", "觉": "覺", "说": "說", "为": "為", "无": "無", "与": "與",
    "体": "體", "万": "萬", "来": "來", "实": "實", "义": "義", "经": "經",
    "论": "論", "号": "號", "当": "當", "从": "從", "学": "學", "断": "斷",
    "边": "邊", "转": "轉", "显": "顯", "现": "現", "应": "應", "处": "處",
    "随": "隨", "点": "點", "师": "師", "净": "淨", "烦": "煩", "恼": "惱",
    "萨": "薩", "刹": "剎", "弥": "彌", "广": "廣", "严": "嚴", "华": "華",
    "会": "會", "个": "個", "于": "於", "么": "麼",      "观": "觀", "释": "釋", "灵": "靈", "杂": "雜", "赖": "賴",
    "龙": "龍", "龟": "龜", "齿": "齒", "关": "關", "两": "兩", "国": "國",
    "气": "氣", "乐": "樂", "恶": "惡", "旧": "舊", "尽": "盡", "独": "獨",
    "变": "變", "数": "數", "归": "歸", "证": "證", "称": "稱", "药": "藥",
    "礼": "禮", "发": "發", "齐": "齊", "黄": "黃", "声": "聲", "读": "讀",
    "铁": "鐵", "禅": "禪", "静": "靜", "摄": "攝", "继": "繼",
    "纤": "纖", "聪": "聰", "肃": "肅", "轻": "輕", "驱": "驅", "笼": "籠",
    "荣": "榮", "庄": "莊", "剑": "劍", "战": "戰", "带": "帶", "壊": "壞",
    # ④ 저본 안에서 실제로 흔들리는 글자
    #    (같은 말뭉치 안에 두 자형이 함께 쓰여, 접지 않으면 검색이 갈린다)
    "徧": "遍", "疎": "疏", "踈": "疏", "薰": "熏", "竝": "並", "歎": "嘆",
    "廻": "迴", "訶": "呵", "葢": "蓋", "盖": "蓋", "祗": "祇", "註": "注",
    "灯": "燈", "弃": "棄", "竪": "豎", "甞": "嘗", "蹟": "跡", "迹": "跡",
    "詶": "酬", "酧": "酬", "慜": "愍", "峯": "峰", "凉": "涼", "燄": "焰",
    "昬": "昏", "煖": "暖", "覔": "覓", "怜": "憐", "皃": "貌", "遶": "繞",
    "瑠": "琉", "璢": "琉", "舩": "船", "巖": "岩", "醎": "鹹", "餝": "飾",
    "罸": "罰", "綵": "彩", "輙": "輒", "鍊": "煉", "谿": "溪", "閙": "鬧",
    "鎻": "鎖", "諠": "喧", "踪": "蹤", "陜": "陝", "賸": "剩", "旛": "幡",
    "冲": "沖", "踴": "踊", "脇": "脅", "飡": "餐", "遯": "遁", "亙": "亘",
    "栢": "柏", "髴": "彿", "鵞": "鵝", "賖": "賒", "貍": "狸", "頺": "頹",
    "隷": "隸", "麪": "麵", "舘": "館", "廼": "迺", "嶋": "島", "鈆": "鉛",
    "阯": "址", "笋": "筍", "騐": "驗", "験": "驗", "躰": "體", "荅": "答",
    "摠": "總", "麽": "麼", "碍": "礙", "虗": "虛", "覈": "核", "讎": "讐",
    "棱": "稜", "賛": "贊",

    # ⑤ 저본에는 한 자형만 있으나 다른 판본·입력기에서 흔히 쓰는 짝
    "縁": "緣", "悩": "惱", "蔵": "藏", "脱": "脫", "蕐": "華", "逹": "達",
    "贒": "賢", "舎": "舍", "繋": "繫", "徴": "徵", "没": "沒", "縦": "縱",
    "莭": "節", "歳": "歲", "睠": "眷", "医": "醫", "銕": "鐵", "渉": "涉",
    "悦": "悅", "甁": "瓶", "歩": "步", "毎": "每", "鼔": "鼓", "兎": "兔",
    "冩": "寫", "渇": "渴", "燐": "憐", "蘓": "蘇", "邨": "村", "売": "賣",
    "担": "擔", "潅": "灌", "謌": "歌", "砕": "碎", "検": "檢", "窃": "竊",
    "鄕": "鄉", "髄": "髓", "涙": "淚", "疣": "肬", "稲": "稻", "禄": "祿",
    "莵": "菟", "緑": "綠", "殻": "殼", "窗": "窓", "裵": "裴", "挿": "插",
    "聨": "聯", "翛": "倏", "鬂": "鬢", "讁": "謫", "齢": "齡", "韈": "襪",
    "顋": "腮", "貮": "貳", "聟": "壻", "蛮": "蠻", "畞": "畝", "籖": "籤",
    "湼": "涅", "饍": "膳",
}


def norm_search(s: str) -> str:
    s = unicodedata.normalize("NFKC", s)
    s = RE_MARKER.sub("", s)
    s = RE_APPARATUS.sub("", s)
    s = "".join(VARIANTS.get(c, c) for c in s)
    s = re.sub(r"[\s。，、．・？！：；「」『』（）()〔〕【】\[\]“”‘’·…—　]", "", s)
    return s


def head_key(s: str, n: int = 14) -> str:
    """정렬용 앞머리 지문."""
    return norm_search(s)[:n]


# ── TXT 파서 ─────────────────────────────────────────────────────────
def parse_txt(path: Path):
    raw = path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp949", "utf-16"):
        try:
            text = raw.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError(f"인코딩 판별 실패: {path}")

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]

    units, order = [], 0
    for b in blocks:
        # 한 블록 안에 위치표지가 여러 개면 표지 단위로 다시 쪼갠다
        parts = re.split(r"(?=\[\d{3,4}[abc]\d{2}\])", b)
        parts = [p.strip() for p in parts if p.strip()]
        for p in parts:
            m = RE_MARKER.match(p)
            order += 1
            units.append({
                "i": order,
                "m": m.group(1) if m else None,
                "cn": re.sub(r"[ \t]+", " ", p).strip(),
            })
    return units


# ── DOCX 파서 ────────────────────────────────────────────────────────
def classify_head(text: str, style: str) -> str:
    """표제를 apparatus / unit / label / structure 로 가른다."""
    t = text.strip()
    if RE_NOTE_HEAD.search(t):
        return "note"
    if t in LABEL_HEADS or label_key(t) in LABEL_HEADS:
        return "label"
    if RE_UNIT_HEAD.match(t):
        return "unit"
    if RE_APPARATUS_HEAD.search(t) or RE_DOC_OUTLINE.match(t):
        return "apparatus"
    return "structure"


def head_level(style: str) -> int:
    m = re.search(r"(\d+)", style or "")
    return int(m.group(1)) if m else 2


def parse_source_docx(path: Path):
    """원문만 담긴 docx(회본 등)를 원문 단위로 읽는다.
    한문 단락만 취하고, 편집 범례·과단 표제 같은 한국어 줄은 버린다."""
    if docx is None:
        raise RuntimeError("python-docx 가 필요합니다: pip install python-docx")
    d = open_docx(path)
    units, order, cur_marker = [], 0, None
    for p in d.paragraphs:
        txt = re.sub(r"[ \t]+", " ", p.text).strip()
        if not txt or not is_source_line(txt):
            continue
        # 【科層】·【科文】·【本文】·【註】 같은 편집 표지로 시작하는 줄은 구조 표시이지
        # 대장경 본문이 아니므로 원문 단위로 세지 않는다
        if re.match(r"^【(?:科層|科文|本文|註|論|疏)】", txt):
            continue
        pieces = [x.strip() for x in
                  re.split(r"(?=\[\d{3,4}[abc]\d{2}\])", txt) if x.strip()]
        for piece in pieces:
            pm = RE_MARKER.match(piece)
            if pm:
                cur_marker = pm.group(1)
            order += 1
            units.append({"i": order, "m": pm.group(1) if pm else cur_marker,
                          "cn": piece})
    return units


PURE_MARKER = re.compile(r"^\[(\d{3,4}[abc]\d{2})\]$")


# ── docx 문단 펴기 ──────────────────────────────────────────────
# 번역 docx 가운데는 한 문단 안에 줄바꿈만으로
#   [0377a05] ⏎ 원문 ⏎ <한문> ⏎ 한국어 번역 ⏎ <번역>
# 을 몰아 담는 형식이 있다. 이대로는 원문과 번역이 한 덩어리로 붙어
# 짝이 지어지지 않으므로, 줄마다 따로 선 문단인 것처럼 펴 준다.
INNER_LABEL = re.compile(r"^(?:원문|한국어(?:\s*(?:직역|번역))?|번역|직역)$")


class _Para:
    """python-docx 문단과 같은 모양(text·style.name)을 가진 대역."""

    __slots__ = ("text", "style")

    def __init__(self, text, style_name):
        self.text = text
        self.style = type("S", (), {"name": style_name})()


class _Doc:
    """문단만 펴 놓은 docx 대역. tables 등은 원본을 그대로 넘긴다."""

    def __init__(self, doc, paras):
        self._doc = doc
        self.paragraphs = paras

    def __getattr__(self, name):
        return getattr(self._doc, name)


def open_docx(path: Path):
    d = docx.Document(str(path))
    out, split_any = [], False
    for p in d.paragraphs:
        lines = [x.strip() for x in p.text.split("\n")]
        lines = [x for x in lines if x]
        # 안쪽에 '원문'·'한국어 번역' 같은 칸 이름이 줄로 서 있을 때만 편다
        if len(lines) > 1 and any(INNER_LABEL.match(x) for x in lines):
            split_any = True
            style = (p.style.name or "").strip()
            for x in lines:
                if INNER_LABEL.match(x):
                    continue
                out.append(_Para(x, style))
        else:
            out.append(p)
    return _Doc(d, out) if split_any else d


def is_translation_only(path: Path) -> bool:
    """원문 없이 번역만 담긴 docx 인가.

    한문 원문 단락이 사실상 없고, '[0459a11]' 처럼 표지만 홀로 선 문단이
    여럿이면 번역 전용 형식으로 본다. 이런 문서는 sources/<id>/원문*.txt
    쪽에서 원문을 따로 대야 하며, merge() 의 표지 대응만으로 짝짓는다."""
    d = open_docx(path)
    paras = [re.sub(r"[ \t]+", " ", p.text).strip() for p in d.paragraphs]
    paras = [t for t in paras if t]
    if not paras:
        return False
    n_src = sum(1 for t in paras if is_source_line(t))
    n_anchor = sum(1 for t in paras if PURE_MARKER.match(t))
    return n_anchor >= 20 and n_src < len(paras) * 0.03


def parse_translation_only_docx(path: Path):
    """표지 단독 문단을 앵커로 삼아 번역만 파싱한다(원문 문단 없음).

    문서는 앞부분(해제) → [표지] 번역… [표지] 번역… → 뒷부분(검증 보고 등)
    순서로 구성된다고 본다. 표지가 처음 나오는 순간부터 '본문'으로 보고,
    표지가 나온 뒤에 다시 나오는 1단계 표제(Heading 1)는 본문이 끝나고
    뒷부분(부록)이 시작된 것으로 본다."""
    d = open_docx(path)
    units, front, appendix = [], [], []
    mode = "front"          # front | body | back
    cur = None

    for p in d.paragraphs:
        txt = re.sub(r"[ \t]+", " ", p.text).strip()
        if not txt:
            continue
        style = (p.style.name or "").strip()

        m = PURE_MARKER.match(txt)
        if m:
            if mode == "back":          # 뒷부분 이후에 표지가 다시 나올 리 없지만 방어적으로
                appendix.append(txt)
                continue
            if cur:
                units.append(cur)
            cur = {"m": m.group(1), "cn": [], "ko": [], "nt": []}
            mode = "body"
            continue

        if mode == "body" and style.lower().startswith("heading"):
            if cur:
                units.append(cur)
                cur = None
            mode = "back"

        if mode == "front":
            front.append(txt)
        elif mode == "back":
            appendix.append(txt)
        elif cur is not None:
            is_note = ("audit" in style.lower()) or ("검증" in style) or ("교감" in style)
            (cur["nt"] if is_note else cur["ko"]).append(txt)

    if cur:
        units.append(cur)

    return {"units": units, "tables": [], "front": front,
            "appendix": appendix, "sections": []}


def is_table_aligned(path: Path) -> bool:
    """대조 본문이 표(위치표지 | 원문 | 번역)로 짜인 docx 인가."""
    if docx is None:
        return False
    d = open_docx(path)
    for t in d.tables:
        if len(t.columns) < 3 or len(t.rows) < 3:
            continue
        head = [c.text.strip() for c in t.rows[0].cells]
        joined = " ".join(head)
        # 번역 칸이 없으면 본문 대조표가 아니다.
        # (‘위치표지 | 원문 | 유형 | 처리’ 같은 교감 목록을 본문으로 오인하지 않도록)
        has_ko = any(k in joined for k in ("번역", "한국어", "직역", "국역"))
        if "위치표지" in joined and "원문" in joined and has_ko:
            return True
        # 표제가 없어도 첫 칸이 위치표지 꼴이고, 어느 칸엔가
        # 한국어 본문이 들어 있으면 대조표로 본다
        if PURE_MARKER.match(head[0] if head else ""):
            body = " ".join(c.text for r in t.rows[:6] for c in r.cells)
            if hangul_ratio(body) > 0.15:
                return True
    return False


def parse_table_docx(path: Path):
    """표로 짜인 대조본을 읽는다.

    각 행이 (위치표지, 원문, 번역) 세 칸으로 이루어진 표를 본문으로 삼고,
    두 칸짜리 표는 서지·용어표로 넘긴다. 표 바깥 문단은 해제로 모은다."""
    if docx is None:
        raise RuntimeError("python-docx 가 필요합니다: pip install python-docx")
    d = open_docx(path)

    units, tables, front = [], [], []

    for t in d.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        rows = [r for r in rows if any(r)]
        if not rows:
            continue
        head = " ".join(rows[0])
        body_table = (len(rows[0]) >= 3 and
                      (("위치표지" in head and "원문" in head)
                       or PURE_MARKER.match(rows[0][0])))
        if not body_table:
            tables.append(rows)          # 서지·용어표
            continue

        start = 1 if ("위치표지" in head or "원문" in head) else 0
        for r in rows[start:]:
            mk_cell, cn_cell = r[0].strip(), r[1].strip()
            ko_cell = r[2].strip() if len(r) > 2 else ""
            m = PURE_MARKER.match(mk_cell) or RE_MARKER.search(mk_cell)
            marker = m.group(1) if m else None
            if not cn_cell and not ko_cell:
                continue
            cn = [x.strip() for x in cn_cell.split("\n") if x.strip()]

            # 번역 칸 안에 '주' 줄이 있고 그 아래 번호 각주가 이어지는 형식이 있다.
            # 줄바꿈으로만 쪼개면 각주가 번역문으로 섞이므로 여기서 갈라낸다.
            ko, nt, in_note = [], [], False
            for x in ko_cell.split("\n"):
                x = x.strip()
                if not x:
                    continue
                if len(x) <= 8 and RE_NOTE_HEAD.search(x):
                    in_note = True
                    continue
                (nt if in_note else ko).append(x)
            units.append({"m": marker, "cn": cn, "ko": ko, "nt": nt})

    for p in d.paragraphs:
        txt = re.sub(r"[ \t]+", " ", p.text).strip()
        if txt and not is_source_line(txt):
            front.append(txt)

    return {"units": units, "tables": tables,
            "front": front, "appendix": [], "sections": []}


# ── KABC(한국불교전서) 번역 docx ───────────────────────────────────
# CBETA 계열과 달리 원문 txt 는 15~16자마다 줄이 끊겨 있어 그대로 실을 수
# 없다. 번역 docx 가 이미 문단으로 이어 붙인 원문을 담고 있으므로,
# 이 갈래는 docx 하나만으로 본문을 세운다.
#
# 블록 차례
#   [저본 위치: 卷上 第2張]  →  원문  →  한국어 직역  →  교감주
RE_KABC_LOC = re.compile(r"^\[저본\s*위치\s*[:：]\s*(.+?)\]\s*$")
RE_KABC_JANG = re.compile(
    r"(卷[上中下一二三四五六七八九十\d]*)?\s*第\s*([一二三四五六七八九十百○\d]+)\s*張")
RE_NOTE_START = re.compile(r"^\[(?:교감|KABC|번역자|위치표지|편집)[^\]]*\]")
# 각주에서 요지로 삼을 줄
RE_NOTE_GIST = re.compile(r"^(?:KABC\s*)?교감\s*내용\s*번역\s*[:：]\s*")
# 각주에서 빼는 줄 — KABC 편집자의 교감문을 그대로 옮긴 대목과,
# 본문을 되풀이하는 대목이다.
RE_NOTE_DROP = re.compile(
    # 줄머리에 '[교감 3)]' 같은 표지가 붙어 있어도 함께 본다
    r"^(?:\[[^\]]*\]\s*)?(?:"
    # ① KABC 편집자의 교감문을 그대로 옮긴 줄
    r"KABC\s*(?:교감\s*원문|저본[·\s]*편집)\s*[:：])|"
    # ② 본문·번역을 되풀이하는 줄 (화면에 이미 있는 것)
    r"^(?:\[[^\]]*\]\s*)?(?:저본|이문)[·\s]*(?:추정)?\s*"
    r"(?:적용[^:：]{0,10}|기준[^:：]{0,12}|번역)\s*[:：]")

HAN_NUM = {"○": 0, "零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
           "六": 6, "七": 7, "八": 8, "九": 9}


def han_to_int(s: str):
    """한자 숫자를 아라비아 숫자로. '第一○張'·'第二十八張' 둘 다 읽는다."""
    s = s.strip()
    if s.isdigit():
        return int(s)
    if not s or any(c not in HAN_NUM and c not in "十百" for c in s):
        return None
    # 자릿수 표기(二十八)와 나열 표기(一○) 를 함께 다룬다
    if "十" in s or "百" in s:
        total, cur = 0, 0
        for c in s:
            if c == "十":
                total += (cur or 1) * 10
                cur = 0
            elif c == "百":
                total += (cur or 1) * 100
                cur = 0
            else:
                cur = HAN_NUM.get(c, 0)
        return total + cur
    n = 0
    for c in s:
        n = n * 10 + HAN_NUM.get(c, 0)
    return n


def kabc_marker(text: str):
    """'[저본 위치: 卷上 第一張]' 의 속살을 '卷上第1張' 꼴로 고른다.

    같은 문헌 안에서도 '第一張' 과 '第2張' 이 섞여 나오므로 숫자를
    아라비아로 통일한다. 표지가 없는 문헌은 None 을 돌려준다."""
    if "없음" in text:
        return None
    m = RE_KABC_JANG.search(text)
    if not m:
        return None
    juan, num = m.group(1) or "", m.group(2)
    n = han_to_int(num)
    if n is None:
        return None
    return f"{juan}第{n}張"


def split_kabc_note(lines):
    """각주 한 덩이를 (요지, 상세) 로 가른다.

    화면에는 요지만 내고, 상세는 손을 얹었을 때 뜨게 한다."""
    keep = [x for x in lines if not RE_NOTE_DROP.match(x)]
    gist = ""
    for x in keep:
        if RE_NOTE_GIST.match(x):
            gist = RE_NOTE_GIST.sub("", x).strip()
            break
    if not gist:
        head = keep[0] if keep else (lines[0] if lines else "")
        gist = re.sub(r"^\[[^\]]*\]\s*", "", head).strip() or head
    # 상세는 손을 얹었을 때 뜨는 쪽지다. '저본 독법'처럼 본문을 통째로
    # 되풀이하는 줄이 있어, 읽을 만한 길이로 잘라 둔다.
    # (원문 전체는 이미 본문 칸에 있으므로 잃는 것이 없다)
    detail = []
    for x in keep:
        x = x.strip()
        if not x or x == gist:
            continue
        if len(x) > 140:
            x = x[:138].rstrip() + "…"
        detail.append(x)
    return gist, detail


def parse_kabc_docx(path: Path):
    """KABC 번역 docx 하나로 본문 단위를 세운다."""
    if docx is None:
        raise RuntimeError("python-docx 가 필요합니다: pip install python-docx")
    d = open_docx(path)
    units, front, cur, marker = [], [], None, None
    note_buf, in_body = [], False

    def flush_note():
        nonlocal note_buf
        if cur is not None and note_buf:
            gist, detail = split_kabc_note(note_buf)
            if gist:
                cur["nt"].append({"t": gist, "d": detail} if detail else gist)
        note_buf = []

    for p in d.paragraphs:
        txt = re.sub(r"[ \t]+", " ", p.text).strip()
        if not txt:
            continue
        style = (p.style.name or "").strip()

        loc = RE_KABC_LOC.match(txt)
        if loc:
            flush_note()
            marker = kabc_marker(loc.group(1))
            in_body = True
            continue
        if style == "Source Text":
            flush_note()
            cur = {"i": len(units) + 1, "m": marker, "cn": [txt],
                   "ko": [], "nt": []}
            units.append(cur)
            in_body = True
            continue
        if style == "Translation Text":
            flush_note()
            if cur is not None:
                cur["ko"].append(txt)
            continue
        if style == "Editorial Note":
            if RE_NOTE_START.match(txt) or not note_buf:
                flush_note()
            note_buf.extend(x.strip() for x in txt.split("\n") if x.strip())
            continue
        if not in_body and style not in LABELS:
            front.append(txt)
    flush_note()

    # 위치표지가 아예 없는 문헌은 단위 순번을 표지로 삼는다
    if units and not any(u["m"] for u in units):
        for u in units:
            u["m"] = None
    return {"units": units, "front": front, "appendix": [], "tables": []}


# ── KABC(한국불교전서) 번역본 ─────────────────────────────────
# KABC 저본은 15~16자마다 줄을 끊어 두어 그대로 실을 수 없다.
# 대신 번역 docx 가 문단으로 묶어 둔 원문을 본문으로 삼는다.
# 문단 차례는 다음과 같다.
#   [저본 위치: 卷上 第1張] / 원문 / <한문> / 한국어 직역 / <번역> / 교감주 / <각주>
KABC_LOC = re.compile(r"^\[저본\s*위치\s*[:：]\s*(.*?)\]\s*$")
KABC_LAYER = re.compile(r"^\[문헌\s*층위\s*[:：]")
KABC_LABELS = {"원문", "한국어 직역", "한국어 번역", "교감주", "기타 각주",
               "주석", "번역", "직역"}
KABC_CONT = re.compile(r"^(?:원문|한국어\s*직역|번역)\s*[(（]\s*이어짐\s*[)）]")
# 각주에서 화면에 늘 보일 줄(요지)과, 얹었을 때만 보일 줄(상세)을 가른다.
NOTE_SUMMARY_KEYS = ("교감 내용 번역", "KABC 교감 내용 번역", "교감 내용 한국어 번역")
CJK_NUM = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7,
           "八": 8, "九": 9, "十": 10, "○": 0, "零": 0}


def cjk_int(s: str):
    """'二八' · '一○' · '三十八' 같은 한자 숫자를 아라비아 숫자로."""
    s = s.strip()
    if s.isdigit():
        return int(s)
    if not s or any(c not in CJK_NUM for c in s):
        return None
    if "十" in s:                       # 三十八 · 十六 · 二十
        a, _, b = s.partition("十")
        return (CJK_NUM[a] if a else 1) * 10 + (CJK_NUM[b] if b else 0)
    if len(s) > 1:                      # 一○ · 二八 처럼 자리마다 적은 꼴
        n = 0
        for c in s:
            n = n * 10 + CJK_NUM[c]
        return n
    return CJK_NUM[s]


def norm_kabc_loc(raw: str) -> str:
    """저본 위치를 '卷上 第1張' 꼴로 고른다."""
    t = re.sub(r"\s+", " ", raw).strip()
    if not t or "없음" in t:
        return ""
    m = re.search(r"(卷[上中下一二三四五六七八九十\d]*)?\s*第\s*"
                  r"([一二三四五六七八九十○\d]+)\s*張", t)
    if not m:
        # '序 — 張 위치표지 이전' 처럼 설명이 붙은 것은 앞머리만 남긴다
        return re.split(r"\s*[—–-]\s*", t)[0].strip()[:12]
    n = cjk_int(m.group(2))
    juan = (m.group(1) or "").strip()
    return f"{juan} 第{n}張".strip() if n is not None else t


def split_kabc_note(text: str):
    """각주 한 덩이를 (요지, 상세)로 가른다."""
    lines = [x.strip() for x in text.split("\n") if x.strip()]
    if not lines:
        return "", ""
    head, summary, detail = lines[0], "", []
    for ln in lines[1:]:
        key, _, val = ln.partition(":")
        if not val:
            key, _, val = ln.partition("：")
        if val and key.strip() in NOTE_SUMMARY_KEYS and not summary:
            summary = val.strip()
        else:
            detail.append(ln)
    if summary:
        tag = re.match(r"^\[[^\]]{1,16}\]", head)
        summary = f"{tag.group(0)} {summary}" if tag else summary
        detail.insert(0, head)
    else:
        summary = head
    return summary, "\n".join(detail)


def read_docx_tables(path: Path):
    """docx 안의 표를 (행 × 칸) 목록으로 읽는다. 서지·용어표 추출용."""
    if docx is None:
        return []
    d = open_docx(path)
    out = []
    for t in d.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if rows:
            out.append(rows)
    return out


def is_kabc_docx(path: Path) -> bool:
    if docx is None:
        return False
    d = open_docx(path)
    for p in d.paragraphs[:400]:
        if KABC_LOC.match(p.text.strip()):
            return True
    return False


def parse_kabc_docx(path: Path):
    """KABC 번역본을 읽는다. 원문·번역·각주를 모두 이 파일에서 얻는다."""
    d = open_docx(path)
    units, front, secs = [], [], []
    loc, cont, started = "", False, False
    note_open = False

    for p in d.paragraphs:
        t = re.sub(r"[ \t]+", " ", p.text).strip()
        if not t:
            continue
        style = (p.style.name or "").strip()

        m = KABC_LOC.match(t)
        if m:
            loc = norm_kabc_loc(m.group(1))
            continue
        if KABC_LAYER.match(t):
            continue
        if KABC_CONT.match(t):
            cont = True
            continue
        if t in KABC_LABELS:
            note_open = t in ("교감주", "기타 각주", "주석")
            continue

        if style == "Source Text":
            if cont and units:
                units[-1]["cn"].append(t)
            else:
                units.append({"i": len(units), "m": loc or None,
                              "cn": [t], "ko": [], "nt": [], "ntd": []})
            cont, started, note_open = False, True, False
            continue
        if style == "Translation Text":
            if units:
                units[-1]["ko"].append(t)
            cont = False
            continue
        if style == "Editorial Note":
            if units:
                # 대괄호 표지로 시작하면 새 각주, 아니면 앞 각주에 이어진다
                if re.match(r"^\[", t) or not units[-1]["nt"]:
                    a, b = split_kabc_note(t)
                    units[-1]["nt"].append(a)
                    units[-1]["ntd"].append(b)
                else:
                    a, b = split_kabc_note(t)
                    joined = "\n".join(x for x in (units[-1]["ntd"][-1], a, b) if x)
                    units[-1]["ntd"][-1] = joined
            continue

        if style.startswith("Heading") and started:
            secs.append({"lv": 2 if style.endswith("2") else 1,
                         "t": t, "at": len(units)})
        elif not started:
            front.append(t)

    # 각주는 여러 문단에 나뉘어 오기도 한다. 덩어리를 다 모은 뒤에
    # 요지와 상세로 갈라, 화면이 바로 쓸 수 있는 꼴로 담는다.
    #   문자열       → 예전처럼 한 줄로 보인다
    #   {t:…, d:[…]} → 요지만 보이고 상세는 얹거나 누르면 편다
    for u in units:
        out = []
        for k in range(len(u["nt"])):
            whole = "\n".join(x for x in (u["nt"][k], u["ntd"][k]) if x)
            gist, detail = split_kabc_note(whole)
            lines = [x.strip() for x in detail.split("\n") if x.strip()]
            out.append({"t": gist, "d": lines} if lines else gist)
        u["nt"] = out
        u.pop("ntd", None)

    # 위치표지가 아예 없는 문헌은 문단 순번을 표지로 삼는다
    if not any(u["m"] for u in units):
        for k, u in enumerate(units, 1):
            u["m"] = f"{k}"
    return {"units": units, "front": front, "secs": secs}


def parse_docx(path: Path):
    if docx is None:
        raise RuntimeError("python-docx 가 필요합니다: pip install python-docx")
    d = open_docx(path)

    # 1) 표: 서지 정보 / 용어 대응표 회수
    tables = []
    for t in d.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        rows = [r for r in rows if any(r)]
        if rows:
            tables.append(rows)

    # 2) 본문 단락 스캔
    blocks, cur_head, cur_marker = [], None, None
    for p in d.paragraphs:
        txt = re.sub(r"[ \t]+", " ", p.text).strip()
        if not txt:
            continue
        style = (p.style.name or "").strip()

        # '주' / '주석' / '각주' / '제1문 주' 는 스타일과 무관하게 각주 시작으로 본다
        if (len(txt) <= 8 and RE_NOTE_HEAD.search(txt)) or (
                RE_NOTE_LABEL_HEAD.match(txt)
                and not RE_NOTE_LABEL_NOT.search(txt)):
            blocks.append({"kind": "head", "hkind": "note", "text": txt,
                           "lv": head_level(style), "m": cur_marker})
            continue
        if txt in LABELS or label_key(txt) in LABELS:
            continue
        # 표지만 홀로 선 문단([0297a11])은 위치 표시일 뿐 본문이 아니다.
        # 걸러 내지 않으면 앞 단위의 번역 끝에 군더더기로 달라붙는다.
        pm = PURE_MARKER.match(txt)
        if pm:
            cur_marker = pm.group(1)
            continue
        if is_unit_label(txt):
            continue

        txt, side = strip_side(txt)
        if not txt:
            continue

        mk = RE_MARKER.search(txt)

        is_head = "heading" in style.lower() or style in (
            "Unit Heading", "Unit Label", "Title", "Subtitle", "Small Meta"
        ) or re.match(r"^(문단\s*\d+|제\s*\d+\s*항|권\s*제?\d+)", txt)

        if is_head:
            kind = classify_head(txt, style)
            if kind == "label":
                continue
            # 한문 표제 바로 뒤에 그 한국어 번역이 오는 문서가 있다.
            # 따로 두면 번역이 사라지므로 앞 표제에 이어 붙인다.
            if (kind == "structure" and blocks
                    and blocks[-1].get("kind") == "head"
                    and blocks[-1].get("hkind") == "structure"
                    and is_source_line(blocks[-1]["text"])
                    and not is_source_line(txt)
                    and " — " not in blocks[-1]["text"]):
                blocks[-1]["text"] += " — " + txt
                cur_head = blocks[-1]["text"]
                continue
            if kind != "apparatus":
                cur_head = txt
                if mk:
                    cur_marker = mk.group(1)
            blocks.append({"kind": "head", "hkind": kind, "text": txt,
                           "lv": head_level(style), "m": cur_marker})
            continue

        if side == "cn" or (side is None and is_source_line(txt)):
            # 한 문단 안에 위치표지가 여럿이면 표지마다 쪼갠다.
            # (docx 가 원문 두 대목을 줄바꿈으로 한 문단에 담는 경우가 있다)
            pieces = [x.strip() for x in
                      re.split(r"(?=\[\d{3,4}[abc]\d{2}\])", txt) if x.strip()]
            for piece in pieces:
                pm = RE_MARKER.match(piece)
                if pm:
                    cur_marker = pm.group(1)
                blocks.append({"kind": "cn", "text": piece,
                               "m": cur_marker, "head": cur_head})
            continue

        kind = "note" if (
            side != "ko" and (
                "note" in style.lower() or "각주" in style or "Editorial" in style
                or NOTE_PREFIX.match(txt))
        ) else "ko"
        blocks.append({"kind": kind, "text": txt, "m": cur_marker, "head": cur_head})

    # 3) 단위 조립
    #    suppress=True 인 동안의 한국어 문단은 '해제'로 보내고 본문에 붙이지 않는다.
    #    한문 단락이 나오면 suppress 는 자동으로 풀린다.
    units, cur = [], None
    front, appendix = [], []
    body_started = False
    suppress = True          # 첫 한문 단락 전까지는 전부 해제
    note_mode = False        # '주' 표제 아래 — 이하 문단은 각주로 모은다
    sections = []            # [{lv, t, i}]  i = 시작 단위 인덱스
    pending_sec = []

    def stash(text):
        (appendix if body_started else front).append(text)

    for b in blocks:
        if b["kind"] == "head":
            if b["hkind"] == "note":
                note_mode = True
                continue
            if note_mode and b["lv"] >= 3:
                # 주석 묶음 안의 소표제 — 절이 아니라 각주 줄로 둔다
                if cur is not None:
                    cur["nt"].append(b["text"])
                continue
            note_mode = False
            if b["hkind"] == "apparatus":
                if cur:
                    cur["sealed"] = True
                suppress = True
                continue
            if b["hkind"] == "structure":
                if not pending_sec or pending_sec[-1]["t"] != b["text"]:
                    pending_sec.append({"lv": b["lv"], "t": b["text"]})
            if cur is not None:
                cur["sealed"] = True
            continue

        if b["kind"] == "cn":
            suppress = False
            note_mode = False
            if not body_started and b["m"]:
                body_started = True
            if cur and not cur["ko"] and not cur["nt"] and not cur.get("sealed"):
                cur["cn"].append(b["text"])
                if cur["m"] is None:
                    cur["m"] = b["m"]
                continue
            if cur:
                units.append(cur)
            for sec in pending_sec:
                sections.append({**sec, "i": len(units)})
            pending_sec = []
            cur = {"m": b["m"], "cn": [b["text"]], "ko": [], "nt": [],
                   "h": b.get("head")}
            continue

        # 한국어 / 주석
        if suppress or cur is None:
            stash(b["text"])
            continue
        if b["kind"] == "ko" and not note_mode:
            cur["ko"].append(b["text"])
        else:
            cur["nt"].append(b["text"])

    if cur:
        units.append(cur)

    # 원문 자리에 한문이 아니라 한국어 작업 설명이 들어앉은 단위는 본문이 아니다.
    # (번역 docx 말미의 「정리 기준」 같은 대목이 '원문' 라벨 아래 놓이는 일이 있다)
    kept = []
    for u in units:
        if any(RE_CJK.search(x) and hangul_ratio(x) < 0.3 for x in u["cn"]):
            kept.append(u)
        else:
            appendix.extend(u["cn"] + u["ko"] + u["nt"])
    units = kept

    # 번역 표지대로 단위를 쪼개면 번호가 밀리므로, 절 표제의 위치도 함께 옮긴다
    units, remap = split_by_ko_markers(units)
    sections = [{**sec, "i": remap.get(sec["i"], sec["i"])} for sec in sections]

    return {"units": units, "tables": tables,
            "front": front, "appendix": appendix, "sections": sections}


def split_by_ko_markers(units):
    """번역 문단마다 위치표지가 붙어 있으면 그 표지대로 원문과 번역을 함께 나눈다.

    원문 수십 줄을 한 '대조 단위'로 묶어 번역한 문서라도, 번역 문단이
    [0091a06] 같은 표지를 달고 있으면 그 표지에서 다음 표지 직전까지의
    원문 줄을 그 번역 옆에 붙일 수 있다. 이렇게 해야 대조가 줄 단위로 맞는다.
    표지가 없는 번역은 앞 조각에 이어 둔다.
    쪼개기 전후의 번호 대응표(remap)를 함께 돌려준다."""
    out, remap = [], {}
    for old, u in enumerate(units):
        remap[old] = len(out)

        marked = [(RE_MARKER.match(k), k) for k in u["ko"]]
        anchors = [m for m, _ in marked if m]
        if len(anchors) < 2 or len(anchors) < len(u["ko"]) * 0.8:
            out.append(u)
            continue

        # 표지별로 번역을 묶는다
        groups = []
        for m, k in marked:
            if m:
                groups.append([m.group(1), [k]])
            elif groups:
                groups[-1][1].append(k)

        # 각 표지가 원문 어느 줄에서 시작하는지 찾는다
        cn_marks = [RE_MARKER.match(c) for c in u["cn"]]
        n, cursor = len(u["cn"]), 0
        starts = []
        for mk, _ in groups:
            pos = None
            for j in range(cursor, n):
                mm = cn_marks[j]
                if mm and mm.group(1) == mk:
                    pos = j
                    break
            starts.append(cursor if pos is None else pos)
            cursor = starts[-1]

        for gi, (mk, kos) in enumerate(groups):
            s0 = starts[gi]
            s1 = starts[gi + 1] if gi + 1 < len(groups) else n
            if s1 < s0:
                s1 = s0
            cn = u["cn"][s0:s1]
            if gi == 0 and s0 > 0:
                cn = u["cn"][:s0] + cn          # 표제·찬자 등 앞머리는 첫 조각에
            rec = {"m": mk, "cn": cn, "ko": kos, "nt": []}
            if gi == 0 and u.get("h"):
                rec["h"] = u["h"]
            out.append(rec)
        out[-1]["nt"] = u["nt"]                 # 각주는 마지막 조각에
    remap[len(units)] = len(out)
    return out, remap


# ── 정렬(병합) ───────────────────────────────────────────────────────
# ── 각주 줄 나누기 ───────────────────────────────────────────────
# 번역 docx 가운데는 그 단위의 교감·판독 메모를 한 문단에 몰아 쓰고
# 「 / 」로만 항목을 가르는 형식이 있다. 화면에서는 한 덩어리로 뭉쳐
# 보이므로, 항목마다 줄을 나눠 둔다.
RE_NOTE_LABEL = re.compile(
    r"^\s*(?:[^\[\]\n]{0,24}?(?:메모|주기|비고))\s*[|｜:：]\s*")
QUOTE_PAIRS = {"\u2018": "\u2019", "\u201c": "\u201d",
               "\u300c": "\u300d", "\u300e": "\u300f",
               "(": ")", "\uff08": "\uff09"}
QUOTE_CLOSE = set(QUOTE_PAIRS.values())


def split_note_line(t: str):
    """각주 한 문단을 「 / 」 자리에서 항목마다 끊는다.

    따옴표·괄호 안의 빗금은 건드리지 않는다. 번역문을 인용하며
    ‘첫째 … / 둘째 …’처럼 쓴 자리가 있어서, 이것까지 끊으면
    한 각주가 엉뚱하게 두 줄로 갈라진다."""
    out, buf, stack, i = [], [], [], 0
    while i < len(t):
        ch = t[i]
        if ch in QUOTE_PAIRS:
            stack.append(QUOTE_PAIRS[ch])
        elif stack and ch == stack[-1]:
            stack.pop()
        elif ch in QUOTE_CLOSE and ch in stack:
            while stack and stack.pop() != ch:
                pass
        if (not stack and ch == "/" and i > 0
                and t[i - 1] in " \t" and i + 1 < len(t) and t[i + 1] in " \t"):
            out.append("".join(buf).strip())
            buf = []
            i += 1
            continue
        buf.append(ch)
        i += 1
    out.append("".join(buf).strip())
    return [x for x in out if x]


def split_note_items(units):
    """각주 문단에서 칸 이름(‘교감·번역 메모 |’)을 떼고 항목마다 줄을 나눈다."""
    for u in units:
        if not u.get("nt"):
            continue
        out = []
        for t in u["nt"]:
            body = RE_NOTE_LABEL.sub("", t, count=1).strip()
            if not body:                      # 칸 이름뿐인 줄은 버린다
                continue
            out.extend(split_note_line(body))
        u["nt"] = out
    return units


def merge(txt_units, dx):
    """원문 TXT 를 정본으로 두고, 번역 docx 단위를 위치표지·본문 대조로 붙인다.

    대응 우선순위
      1) 위치표지 일치
      2) 한문 앞머리 지문 일치
      3) 원문 전체를 한 줄로 이은 뒤 부분문자열 탐색 (docx 가 원문을 잘게 쪼갠 경우)
    여러 docx 단위가 한 TXT 단위에 대응하면 번역을 순서대로 이어 붙인다.
    """
    dx_units = dx["units"] if dx else []

    # docx 가 원문을 사실상 전부 담고 있으면 docx 의 세분 단위를 정본 순서로 삼는다
    if txt_units and dx_units:
        t_all = norm_search(" ".join(u["cn"] for u in txt_units))
        d_all = norm_search(" ".join(" ".join(u["cn"]) for u in dx_units))
        if len(t_all) and len(d_all) / len(t_all) >= 0.93 and len(dx_units) > len(txt_units):
            txt_units = []

    # docx 단독 문헌
    if not txt_units:
        out = []
        for n, u in enumerate(dx_units, 1):
            rec = {"i": n, "m": u["m"], "cn": u["cn"], "ko": u["ko"], "nt": u["nt"]}
            if u.get("h"):
                rec["h"] = u["h"]
            out.append(rec)
        return out, {k: k for k in range(len(dx_units))}

    out = []
    for tu in txt_units:
        rec = {"i": tu["i"], "m": tu["m"], "cn": [tu["cn"]], "ko": [], "nt": []}
        out.append(rec)

    # 한 위치표지에 원문 단락이 여럿 딸리는 일이 흔하다(問·答이 같은 행에서 시작하는 등).
    # 그래서 표지·앞머리 모두 '대기열'로 두고, 한 번 짝지어진 단위는 다시 쓰지 않는다.
    by_marker, by_head = {}, {}
    flat, offsets = [], []
    pos = 0
    for idx, rec in enumerate(out):
        if rec["m"]:
            by_marker.setdefault(rec["m"], []).append(idx)
        by_head.setdefault(head_key(rec["cn"][0]), []).append(idx)
        n = norm_search(rec["cn"][0])
        flat.append(n)
        offsets.append((pos, idx))
        pos += len(n)
    haystack = "".join(flat)
    starts = [o[0] for o in offsets]

    import bisect

    def locate(needle):
        if len(needle) < 6:
            return None
        p = haystack.find(needle)
        if p < 0:
            return None
        return offsets[bisect.bisect_right(starts, p) - 1][1]

    # docx 한 덩어리가 원문 여러 단락에 걸치는 경우가 흔하다.
    # 조각마다 대응 위치를 찾아, 걸친 범위를 한 단위로 합친다.
    taken = set()

    def pick(queue):
        for i in queue:
            if i not in taken:
                return i
        return None

    def resolve(piece):
        # 본문 앞머리가 가장 확실한 단서이므로 먼저 본다
        k = head_key(piece)
        if k and k in by_head:
            i = pick(by_head[k])
            if i is not None:
                return i
        mk = RE_MARKER.search(piece)
        if mk and mk.group(1) in by_marker:
            i = pick(by_marker[mk.group(1)])
            if i is not None:
                return i
        return locate(norm_search(piece)[:24])

    unmatched, dxmap, spans = [], {}, []
    owners = set()
    for dxi, u in enumerate(dx_units):
        if u["cn"]:
            targets = [t for t in (resolve(c) for c in u["cn"]) if t is not None]
        else:
            # 원문 조각 없이 표지만 있는 번역 닻.
            # 원문 한 줄이 길면 그 줄 하나에 번역 문단이 여럿 달린다.
            # 그러므로 이미 다른 번역이 붙은 자리라도 이어 붙일 수 있어야 한다.
            targets = []
            if u["m"] and u["m"] in by_marker:
                q = by_marker[u["m"]]
                i = pick(q)                  # 아직 비어 있는 자리가 있으면 그쪽
                if i is None:
                    i = q[0]                 # 없으면 같은 표지의 첫 자리에 이어 붙인다
                targets = [i]
        if not targets:
            unmatched.append((dxi, u))
            continue
        # docx 한 단위가 담은 원문 조각 수보다 훨씬 넓은 범위에 걸쳐 있다면
        # 어느 한 조각이 엉뚱한 자리에 붙은 것이다. 그 이상치는 버린다.
        targets.sort()
        owner = targets[0]
        span_max = len(u["cn"]) * 2 + 10
        targets = [t for t in targets if t - owner <= span_max]
        last = targets[-1]
        taken.update(targets)
        dxmap[dxi] = owner
        owners.add(owner)
        spans.append((owner, last))
        out[owner]["ko"].extend(u["ko"])
        out[owner]["nt"].extend(u["nt"])
        if u.get("h") and "h" not in out[owner]:
            out[owner]["h"] = u["h"]

    # 다른 번역이 걸려 있지 않은 중간 단락만 앞 단위로 흡수한다
    absorbed = {}
    for owner, last in spans:
        for k in range(owner + 1, last + 1):
            if k not in owners and k not in absorbed:
                absorbed[k] = owner

    # 묶음형 docx 대응
    #   원문 여러 줄(수십 줄)을 한 '대조 단위'로 묶어 번역한 문서에서는
    #   단위 사이에 낀 원문 줄이 어디에도 걸리지 않고 남는다.
    #   이때는 docx 단위를 구간 경계로 삼아, 다음 경계 직전까지를 그 단위에 붙인다.
    if dx_units:
        per_unit = sorted(len(u["cn"]) for u in dx_units if u["cn"])
        if not per_unit:
            per_unit = [0]
        median_cn = per_unit[len(per_unit) // 2]
        if median_cn >= 5:                       # 한 단위가 원문 5줄 이상을 묶는 문서
            bounds = sorted(owners)
            for n, start in enumerate(bounds):
                stop = bounds[n + 1] if n + 1 < len(bounds) else len(out)
                for k in range(start + 1, stop):
                    if k not in owners and k not in absorbed:
                        absorbed[k] = start
    for k in sorted(absorbed):
        out[absorbed[k]]["cn"].extend(out[k]["cn"])

    # 원문 TXT 에 없는 번역 단위(다른 저본·회본 등)는 별권으로 이어 붙인다
    base = len(out)
    for n, (dxi, u) in enumerate(unmatched, 1):
        dxmap[dxi] = len(out)
        out.append({"i": base + n, "m": u["m"], "cn": u["cn"],
                    "ko": u["ko"], "nt": u["nt"], "x": 1})

    final, remap = [], {}
    for idx, rec in enumerate(out):
        if idx in absorbed:
            continue
        remap[idx] = len(final)
        rec["i"] = len(final) + 1
        final.append(rec)
    dxmap = {k: remap.get(v, remap.get(absorbed.get(v, v), 0))
             for k, v in dxmap.items()}
    return final, dxmap


# ── 표에서 용어표 추출 ───────────────────────────────────────────────
def pick_glossary(tables):
    gloss = []
    for rows in tables:
        header = [h.strip() for h in rows[0]]
        joined = " ".join(header)
        if ("원어" in joined or "번역어" in joined or "원문" in joined) and len(header) >= 2:
            hi = 0
            for i, h in enumerate(header):
                if "원어" in h or h == "원문":
                    hi = i
            ki = 1 if hi == 0 else 0
            for i, h in enumerate(header):
                if "번역" in h and i != hi:
                    ki = i
            for r in rows[1:]:
                if len(r) > max(hi, ki) and r[hi] and r[ki]:
                    gloss.append({"cn": r[hi], "ko": r[ki],
                                  "note": r[max(hi, ki) + 1] if len(r) > max(hi, ki) + 1 else ""})
            if gloss:
                return gloss
    return gloss


def pick_biblio(tables):
    bib = {}
    for rows in tables:
        if len(rows[0]) == 2 and rows[0][0] in ("항목", "구분") or (
            len(rows[0]) == 2 and any("문헌명" in r[0] for r in rows)
        ):
            for r in rows:
                if len(r) >= 2 and r[0] and r[1] and r[0] not in ("항목", "구분"):
                    bib.setdefault(r[0], r[1])
    return bib


# ── 실행 ─────────────────────────────────────────────────────────────
def detect_juan_from_source(units):
    """원문 안에 남아 있는 권 표제(…卷第一 / …卷上)로 분권을 잡는다."""
    secs, seen = [], set()
    for idx, u in enumerate(units):
        for line in u["cn"]:
            line = line.strip()
            # 교감 표지와 서명 부분은 빼고 '卷第三' 꼴만 견준다.
            # 같은 권의 여는 줄과 닫는 줄, 서명에 오자가 섞인 줄을
            # 서로 다른 권으로 세지 않기 위함이다.
            key = RE_APPARATUS.sub("", line)
            key = key[key.find("卷"):] if "卷" in key else key
            if RE_JUAN_LINE.match(line) and key not in seen:
                seen.add(key)
                secs.append({"lv": 1, "t": line, "i": idx})
                break
    return secs


def chunk_by_chars(units, start, end, target=26000):
    """단위 수가 적어도 글자 수가 많으면 쪽 경계에서 끊는다.
    묶음형 번역처럼 한 단위가 원문 수십 줄을 담는 문헌용."""
    out, acc, last_page = [], 0, None
    for i in range(start, end):
        u = units[i]
        page = (u.get("m") or "")[:-3]
        if acc >= target and page and page != last_page:
            out.append(i)
            acc = 0
        acc += len(norm_search(" ".join(u["cn"])))
        if page:
            last_page = page
    if out and (end - out[-1]) < 3:
        out.pop()
    return out


def chunk_by_page(units, start, end, target=90):
    """표제가 없는 구간을 위치표지의 쪽 번호로 끊는다.
    쪽이 바뀌는 자리에서만 자르므로 문맥 한가운데가 갈라지지 않는다."""
    breaks = []
    last_page = None
    for i in range(start, end):
        m = units[i].get("m")
        if not m:
            continue
        page = m[:-3]                      # 1206c22 → 1206
        if last_page is not None and page != last_page:
            breaks.append(i)
        last_page = page
    if not breaks:
        return []

    out, prev = [], start
    for b in breaks:
        if b - prev >= target:
            out.append(b)
            prev = b
    if out and (end - out[-1]) < target // 2:   # 꼬리가 너무 짧으면 앞에 붙인다
        out.pop()
    return out


def label_range(units, a, b):
    """구간의 위치표지 범위를 이름으로 삼는다 — 1206c22–1210b04"""
    ms = [units[i]["m"] for i in range(a, b) if units[i].get("m")]
    if not ms:
        return f"{a + 1}–{b}단위"
    return ms[0] if len(ms) == 1 else f"{ms[0]}–{ms[-1]}"


def subdivide(units, chapters, limit=150):
    """지나치게 긴 권을 쪽 단위로 잘게 나눈다."""
    if not chapters:
        return chapters
    out = []
    for n, c in enumerate(chapters):
        end = chapters[n + 1]["i"] if n + 1 < len(chapters) else len(units)
        out.append(c)
        if end - c["i"] <= limit:
            continue
        cuts = chunk_by_page(units, c["i"], end)
        for k, b in enumerate(cuts):
            stop = cuts[k + 1] if k + 1 < len(cuts) else end
            out.append({"lv": c["lv"], "i": b,
                        "t": f"{c['t']} ({k + 2}) {label_range(units, b, stop)}"})
    out.sort(key=lambda x: x["i"])
    return out


def build_sections(sections, n_units):
    """빈 절을 걷어내고, 화면 분할에 쓸 '권(chapter)' 층위를 고른다."""
    secs = [s for s in sections if 0 <= s["i"] < n_units]
    secs.sort(key=lambda x: (x["i"], x["lv"]))

    # 같은 위치에 겹친 표제는 가장 상위만 남긴다
    dedup, last_i = [], None
    for s in secs:
        if s["i"] == last_i:
            continue
        dedup.append(s)
        last_i = s["i"]
    secs = dedup

    # 내용이 없는 절 제거(다음 절과 시작 위치가 같은 경우)
    kept = []
    for n, s in enumerate(secs):
        nxt = secs[n + 1]["i"] if n + 1 < len(secs) else n_units
        if nxt > s["i"]:
            kept.append(s)
    secs = kept
    if not secs:
        return [], []

    # 짧은 문헌은 굳이 나누지 않는다. 절 표제는 본문 안 소제목으로만 쓴다.
    if n_units < 60:
        return secs, []

    # 권 층위 고르기: 2~120권으로 나뉘고 한 권이 평균 3단위 이상인 가장 상위 층위
    # (48권짜리 『유가론기』처럼 권수가 많은 문헌도 권 표제를 살려 쓰기 위함)
    levels = sorted({s["lv"] for s in secs})
    chapter_lv = None
    for lv in levels:
        n = sum(1 for s in secs if s["lv"] <= lv)
        if 2 <= n <= 120 and n_units / n >= 3:
            chapter_lv = lv
            break
    if chapter_lv is None:
        return secs, []

    chapters = [s for s in secs if s["lv"] <= chapter_lv]
    if chapters and chapters[0]["i"] > 0:
        chapters.insert(0, {"lv": chapter_lv, "t": "권두", "i": 0})
    return secs, chapters


def build_work(entry):
    wid = entry["id"]
    wdir = SRC / wid
    txt_units, dx = [], None

    # 원문은 여러 개일 수 있다. 회본(본문 + 주석서)이 그런 경우로,
    #   원문-1-십이문론.txt / 원문-2-종치의기.txt
    # 처럼 이름을 붙이면 파일명 순서대로 이어 붙인다.
    tps = sorted(list(wdir.glob("원문*.txt")) + list(wdir.glob("원문*.docx")))
    for tp in tps:
        part = parse_source_docx(tp) if tp.suffix == ".docx" else parse_txt(tp)
        base = len(txt_units)
        for u in part:
            u["i"] += base
            u["src"] = tp.stem
        txt_units.extend(part)
    dp = wdir / "번역.docx"
    kabc = entry.get("source") == "KABC"
    if kabc:
        # 한국불교전서 계열은 저본 txt 가 15~16자마다 끊겨 있어 본문으로
        # 쓸 수 없다. 번역 docx 안의 원문이 이미 문단으로 이어져 있으므로
        # 그쪽 하나로 단위를 세운다. (저본 txt 는 대조 근거로만 둔다)
        dx = parse_kabc_docx(dp)
        dx.setdefault("sections", dx.get("secs", []))
        dx.setdefault("appendix", [])
        dx.setdefault("tables", read_docx_tables(dp))
        units, dxmap = dx["units"], {}
    elif dp.exists():
        if is_table_aligned(dp):
            dx = parse_table_docx(dp)
        elif is_translation_only(dp):
            dx = parse_translation_only_docx(dp)
        else:
            dx = parse_docx(dp)

    if not kabc:
        units, dxmap = merge(txt_units, dx)
        units = split_note_items(units)

    # ── 분권 구성 ────────────────────────────────────────────────
    raw_secs = []
    if dx and not kabc:
        for sec in dx["sections"]:
            j = dxmap.get(sec["i"])
            if j is None:
                # 대응 못 찾은 표제는 그 뒤 첫 대응 지점으로 민다
                later = [dxmap[k] for k in sorted(dxmap) if k >= sec["i"]]
                if not later:
                    continue
                j = later[0]
            raw_secs.append({"lv": sec["lv"], "t": sec["t"], "i": j})
    if not raw_secs:
        raw_secs = detect_juan_from_source(units)

    sections, chapters = build_sections(raw_secs, len(units))

    # 번역 docx가 권마다 표제를 달아 두지 않은 문헌은, 권 나눔이 엉뚱하게 잡힌다.
    # (한 권 제목이 뒤 권까지 끌려가 '…(2) (3)'처럼 되풀이된다)
    # 이럴 때는 원문에 남아 있는 권 표제를 따르는 편이 낫다.
    # 번역 docx가 권마다 표제를 달아 두면 그쪽이 한국어 병기까지 있어 낫다.
    # 표제가 거의 없을 때만(=권이 3개 이하로 잡힐 때) 원문 표제로 갈아탄다.
    if len(chapters) <= 3:
        juan = detect_juan_from_source(units)
        if len(juan) >= 4 and len(juan) > len(chapters):
            _, chapters = build_sections(juan, len(units))


    # 단위는 적지만 글자 수가 많은 문헌(묶음형 번역)도 나눈다
    total_chars = sum(len(norm_search(" ".join(u["cn"]))) for u in units)
    if not chapters and total_chars > 40000 and len(units) <= 120:
        cuts = [0] + chunk_by_chars(units, 0, len(units))
        if len(cuts) > 1:
            chapters = [
                {"lv": 1, "i": b,
                 "t": label_range(units, b,
                                  cuts[k + 1] if k + 1 < len(cuts) else len(units))}
                for k, b in enumerate(cuts)]

    # 표제가 없어 권이 안 잡힌 긴 문헌은 위치표지 쪽으로 나눈다
    if not chapters and len(units) > 120:
        cuts = [0] + chunk_by_page(units, 0, len(units))
        if len(cuts) > 1:
            chapters = [
                {"lv": 1, "i": b,
                 "t": label_range(units, b,
                                  cuts[k + 1] if k + 1 < len(cuts) else len(units))}
                for k, b in enumerate(cuts)]

    # 권 하나가 너무 길면 더 잘게
    chapters = subdivide(units, chapters)

    # 각 단위에 소속 권 번호 부여
    if chapters:
        ci, nxt = 0, (chapters[1]["i"] if len(chapters) > 1 else len(units))
        for idx, u in enumerate(units):
            while ci + 1 < len(chapters) and idx >= chapters[ci + 1]["i"]:
                ci += 1
            u["c"] = ci

    # ── 통계 ─────────────────────────────────────────────────────
    n_ko = sum(1 for u in units if u["ko"])
    chars_cn = sum(len(norm_search(" ".join(u["cn"]))) for u in units)
    chars_done = sum(len(norm_search(" ".join(u["cn"]))) for u in units if u["ko"])
    chars_ko = sum(len("".join(u["ko"])) for u in units)
    markers = sorted({u["m"] for u in units if u["m"] and not u.get("x")})

    # 저본에 실린 도판을 그 자리의 단위에 붙인다.
    # registry 의 figures = {"위치표지": "파일명"} 을 따르고,
    # 파일은 assets/figures/<문헌id>/ 에 둔다.
    figs = entry.get("figures") or {}
    if figs:
        first, last = {}, {}
        for u in units:
            first.setdefault(u.get("m"), u)
            last[u.get("m")] = u
        marks = sorted(m for m in first if m)
        for marker, fname in sorted(figs.items()):
            u = first.get(marker)
            if u is None:
                # 저본에서 도판만 놓인 표지는 원문이 비어 단위가 서지 않는다.
                # 그 앞 표지의 마지막 단위 끝에 붙여 저본과 같은 자리에 오게 한다.
                prev = [m for m in marks if m < marker]
                u = last.get(prev[-1]) if prev else None
            if u is None:
                print(f"  ! {wid}: 도판 {fname} 의 자리 [{marker}] 를 찾지 못했습니다")
                continue
            u.setdefault("fig", []).append(fname)
            # 도판을 실제로 띄우므로, 그 자리를 말로 때운 번역 문단
            # ('문자 본문은 없으며, 이 위치에 …가 배치된다')은 군더더기가 된다.
            u["ko"] = [t for t in u["ko"] if not RE_FIG_CAPTION.search(t)]

    meta = dict(entry)
    units_marks = [u["m"] for u in units if u.get("m")]
    meta.pop("figures", None)
    meta.update({
        "units": len(units),
        "translated": n_ko,
        "coverage": round(chars_done / chars_cn, 4) if chars_cn else 0,
        "chars_cn": chars_cn,
        "chars_ko": chars_ko,
        # 위치표지는 대개 사전순이 곧 문헌 순서지만, 한불전의 張 표지는
        # 「第9張」이 「第10張」보다 뒤로 밀리므로 본문에 나온 차례를 따른다.
        # (회본처럼 원문을 이어 붙인 문헌은 사전순 쪽이 낫다)
        "range": ([units_marks[0], units_marks[-1]] if kabc and units_marks
                  else ([markers[0], markers[-1]] if markers else None)),
        "has_source": bool(txt_units),
        "has_translation": bool(dx),
        "chapters": len(chapters),
    })
    if dx:
        gl = pick_glossary(dx["tables"])
        meta["biblio_extracted"] = pick_biblio(dx["tables"])
    else:
        gl = []

    doc = {
        "meta": meta,
        "glossary": gl,
        "front": (dx["front"] + dx["appendix"]) if dx else [],
        "chapters": [{"t": c["t"], "i": c["i"]} for c in chapters],
        "sections": [{"lv": s["lv"], "t": s["t"], "i": s["i"]} for s in sections],
        "units": units,
    }
    (OUT / "works").mkdir(parents=True, exist_ok=True)
    with open(OUT / "works" / f"{wid}.json", "w", encoding="utf-8") as f:
        json.dump(doc, f, ensure_ascii=False, separators=(",", ":"))
    return meta


def discover(reg):
    """sources/ 안에 있으나 registry 에 없는 문헌을 자동 편입한다.
    서지를 몰라도 일단 읽히도록 최소 정보만 채워 둔 뒤,
    registry.json 에 항목을 추가하면 그때부터 정식 서지가 적용된다."""
    known = {w["id"] for w in reg["works"]}
    added = []
    if not SRC.exists():
        return added
    for d in sorted(SRC.iterdir()):
        if not d.is_dir() or d.name in known:
            continue
        if not any((d / f).exists() for f in ("원문.txt", "번역.docx")):
            continue
        added.append({
            "id": d.name,
            "title_cn": d.name,
            "title_ko": d.name,
            "author_cn": "", "author_ko": "미상", "dynasty": "",
            "canon": "서지 미기재", "canon_label": "미분류",
            "canon_id": "Z", "canon_vol": 999, "canon_no": 999,
            "canon_ko": "", "collection": "미분류", "tags": [],
            "verify": True,
        })
    return added


def sync_variant_tables(root: Path) -> None:
    """이체자 표를 화면 쪽 js 에도 그대로 옮겨 적는다.

    검색 색인은 여기서(파이썬), 검색어 정규화는 브라우저(js)에서 이루어진다.
    두 표가 어긋나면 원문에 있는 글자를 찾지 못하므로, 표는 이 파일 하나만
    고치고 나머지는 이 함수가 맞춰 준다."""
    body = ",\n  ".join(
        f"'{k}':'{v}'" for k, v in sorted(VARIANTS.items()) if k != v)
    targets = {"assets/js/search.worker.js": "const VAR = {",
               "assets/js/app.js": "const VAR_FOLD = {"}
    for rel, head in targets.items():
        f = root / rel
        if not f.exists():
            continue
        t = f.read_text(encoding="utf-8")
        i = t.find(head)
        if i < 0:
            print(f"  ! {rel}: 이체자 표를 찾지 못했습니다")
            continue
        j = t.index("};", i) + 2
        new = t[:i] + head + "\n  " + body + "\n};" + t[j:]
        if new != t:
            f.write_text(new, encoding="utf-8")
            print(f"  · {rel} 이체자 표 갱신 ({len(VARIANTS)}자)")


def main():
    reg = json.loads((ROOT / "tools" / "registry.json").read_text(encoding="utf-8"))
    fresh = discover(reg)
    if fresh:
        print("registry 에 없는 문헌을 자동 편입합니다 "
              "(tools/registry.json 에 서지를 채워 주세요):")
        for e in fresh:
            print("   +", e["id"])
        reg["works"] = reg["works"] + fresh
    sync_variant_tables(ROOT)

    metas = []
    for entry in reg["works"]:
        m = build_work(entry)
        metas.append(m)
        print(f"  {m['id']:<22} 단위 {m['units']:>5}  번역 {m['translated']:>5}"
              f"  ({m['coverage']*100:5.1f}%)  {m['range']}")
    # 대장경 순서(T → X → L, 책 번호, 경 번호)로 정렬해 둔다
    CANON_ORDER = {"T": 0, "X": 1, "L": 2, "K": 3, "B": 4, "ZW": 5, "BJ": 6, "HB": 6}
    metas.sort(key=lambda m: (CANON_ORDER.get(m.get("canon_id", ""), 9),
                              m.get("canon_vol", 0), m.get("canon_no", 0)))
    manifest = {
        "generated": reg.get("version", "1"),
        "site": reg.get("site", {}),
        "rights": reg.get("rights", {}),
        "works": [{k: v for k, v in m.items() if k != "biblio_extracted"} for m in metas],
        "totals": {
            "works": len(metas),
            "units": sum(m["units"] for m in metas),
            "translated": sum(m["translated"] for m in metas),
            "chars_cn": sum(m["chars_cn"] for m in metas),
        },
    }
    with open(OUT / "manifest.json", "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=1)
    # 재배포 조건 점검 — 출처마다 조건이 다르므로 빌드할 때마다 확인한다
    srcinfo = reg.get("rights", {}).get("sources", {})
    flagged = {}
    for m in metas:
        src = m.get("source", "CBETA")
        info = srcinfo.get(src, {})
        if info.get("ok") is not True:
            flagged.setdefault(src, []).append(m["id"])
    for src, ids in flagged.items():
        info = srcinfo.get(src, {})
        mark = "재배포 금지" if info.get("ok") is False else "조건 확인 필요"
        print(f"\n[주의] 출처 {src} — {mark}")
        print(f"        {info.get('condition', '이용 조건을 확인하십시오.')}")
        print(f"        해당 문헌: {', '.join(ids)}")

    print(f"\n총 {manifest['totals']['works']}종 · "
          f"{manifest['totals']['units']}단위 · "
          f"원문 {manifest['totals']['chars_cn']:,}자")


if __name__ == "__main__":
    main()
